from io import BytesIO

import pytest

from app.lib.documents.parser import DocumentParseError, parse_uploaded_document
from app.shared.settings import get_settings


def test_parse_markdown_uses_utf8_text_path() -> None:
    parsed = parse_uploaded_document(
        filename="guide.md",
        content_type="text/markdown",
        data="# 标题\n\n这是 UTF-8 Markdown 内容。".encode("utf-8"),
        title="知识文档",
        settings=get_settings(),
    )

    assert parsed.title == "知识文档"
    assert parsed.content == "# 标题\n这是 UTF-8 Markdown 内容。"
    assert parsed.doc_type == "markdown"
    assert parsed.metadata["ocr_used"] is False


def test_parse_text_uses_charset_fallback_for_gb18030() -> None:
    settings = get_settings()
    parsed = parse_uploaded_document(
        filename="legacy.txt",
        content_type="text/plain",
        data="中文旧编码内容".encode("gb18030"),
        title=None,
        settings=settings,
    )

    assert parsed.content == "中文旧编码内容"
    assert parsed.doc_type == "text"
    assert parsed.metadata["parser"] == "gb18030"


def test_parse_docx_preserves_paragraph_table_order() -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("开头段落")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "表格左"
    table.rows[0].cells[1].text = "表格右"
    document.add_paragraph("结尾段落")

    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_uploaded_document(
        filename="ordered.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=buffer.getvalue(),
        title=None,
        settings=get_settings(),
    )

    assert parsed.doc_type == "docx"
    assert parsed.content.splitlines() == ["开头段落", "表格左\t表格右", "结尾段落"]


def test_parse_pdf_uses_pypdf_when_text_is_sufficient(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self):
            return "PDF 正文内容 " * 20

    class FakeReader:
        def __init__(self, _):
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    parsed = parse_uploaded_document(
        filename="text.pdf",
        content_type="application/pdf",
        data=b"%PDF fake",
        title=None,
        settings=get_settings().model_copy(update={"document_ocr_min_chars_per_page": 20}),
    )

    assert parsed.doc_type == "pdf"
    assert "PDF 正文内容" in parsed.content
    assert parsed.metadata["parser"] == "pypdf"
    assert parsed.metadata["ocr_used"] is False


def test_parse_pdf_records_failed_pypdf_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class FailingPage:
        def extract_text(self):
            raise RuntimeError("page stream failed")

    class TextPage:
        def extract_text(self):
            return "PDF 正文内容 " * 20

    class FakeReader:
        def __init__(self, _):
            self.pages = [FailingPage(), TextPage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    parsed = parse_uploaded_document(
        filename="partial.pdf",
        content_type="application/pdf",
        data=b"%PDF fake",
        title=None,
        settings=get_settings().model_copy(update={"document_ocr_min_chars_per_page": 20}),
    )

    assert parsed.doc_type == "pdf"
    assert "PDF 正文内容" in parsed.content
    assert parsed.metadata["pypdf_failed_pages"] == [1]
    assert parsed.metadata["pypdf_failed_page_count"] == 1


def test_parse_pdf_reports_ocr_unavailable_when_text_is_empty(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self):
            return ""

    class FakeReader:
        def __init__(self, _):
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    monkeypatch.setattr("app.lib.documents.parser.shutil.which", lambda _: None)

    with pytest.raises(DocumentParseError) as exc_info:
        parse_uploaded_document(
            filename="scan.pdf",
            content_type="application/pdf",
            data=b"%PDF fake",
            title=None,
            settings=get_settings(),
        )

    assert exc_info.value.code == "OCR_UNAVAILABLE"
    assert "tesseract" in exc_info.value.message


def test_parse_pdf_keeps_partial_text_when_ocr_is_unavailable(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FakePage:
        def extract_text(self):
            return "页眉"

    class FakeReader:
        def __init__(self, _):
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    monkeypatch.setattr("app.lib.documents.parser.shutil.which", lambda _: None)

    parsed = parse_uploaded_document(
        filename="mixed.pdf",
        content_type="application/pdf",
        data=b"%PDF fake",
        title=None,
        settings=get_settings().model_copy(update={"document_ocr_min_chars_per_page": 50}),
    )

    assert parsed.doc_type == "pdf"
    assert parsed.content == "页眉"
    assert parsed.metadata["ocr_fallback_to_pypdf"] is True
    assert parsed.metadata["ocr_error_code"] == "OCR_UNAVAILABLE"


def test_parse_rejects_empty_text_document() -> None:
    with pytest.raises(DocumentParseError) as exc_info:
        parse_uploaded_document(
            filename="empty.txt",
            content_type="text/plain",
            data=b" \n\t ",
            title=None,
            settings=get_settings(),
        )

    assert exc_info.value.code == "DOCUMENT_TEXT_EMPTY"


def test_parse_rejects_file_too_large() -> None:
    with pytest.raises(DocumentParseError) as exc_info:
        parse_uploaded_document(
            filename="large.txt",
            content_type="text/plain",
            data=b"abcd",
            title=None,
            settings=get_settings().model_copy(update={"document_max_file_bytes": 3}),
        )

    assert exc_info.value.code == "FILE_TOO_LARGE"
    assert exc_info.value.status_code == 413


def test_parse_pdf_rejects_too_many_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self):
            return "正文"

    class FakeReader:
        def __init__(self, _):
            self.pages = [FakePage(), FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    with pytest.raises(DocumentParseError) as exc_info:
        parse_uploaded_document(
            filename="too-many.pdf",
            content_type="application/pdf",
            data=b"%PDF fake",
            title=None,
            settings=get_settings().model_copy(update={"document_max_pdf_pages": 1}),
        )

    assert exc_info.value.code == "PDF_TOO_MANY_PAGES"
    assert exc_info.value.status_code == 413


def test_parse_pdf_rejects_ocr_when_page_count_exceeds_ocr_limit(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FakePage:
        def extract_text(self):
            return ""

    class FakeReader:
        def __init__(self, _):
            self.pages = [FakePage(), FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    with pytest.raises(DocumentParseError) as exc_info:
        parse_uploaded_document(
            filename="scan.pdf",
            content_type="application/pdf",
            data=b"%PDF fake",
            title=None,
            settings=get_settings().model_copy(
                update={"document_max_pdf_pages": 10, "document_ocr_max_pages": 1}
            ),
        )

    assert exc_info.value.code == "OCR_TOO_MANY_PAGES"
    assert exc_info.value.status_code == 413
